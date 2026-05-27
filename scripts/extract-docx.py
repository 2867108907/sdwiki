#!/usr/bin/env python3
"""
Extract text from Word docx file for SDWiki ingestion
Usage: ./scripts/extract-docx.py input.docx
"""

import sys
import os

try:
    from docx import Document
except ImportError:
    print("Error: python-docx not installed", file=sys.stderr)
    print("Install with: pip install python-docx", file=sys.stderr)
    sys.exit(1)

def main():
    if len(sys.argv) != 2:
        print("Usage: ./extract-docx.py <input.docx>", file=sys.stderr)
        sys.exit(1)

    input_file = sys.argv[1]

    if not os.path.exists(input_file):
        print(f"Error: Input file {input_file} not found", file=sys.stderr)
        sys.exit(1)

    ext = os.path.splitext(input_file)[1].lower()
    if ext != '.docx':
        print(f"Error: Only .docx files are supported, got {ext}", file=sys.stderr)
        sys.exit(1)

    doc = Document(input_file)
    lines = [para.text for para in doc.paragraphs]

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                lines.append(cell.text)

    print('\n'.join(lines))
    print(f"\n--- Extracted {len(lines)} lines from {input_file} ---", file=sys.stderr)

if __name__ == "__main__":
    main()
